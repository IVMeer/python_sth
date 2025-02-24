import openpyxl
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# 读取 Excel 文件
def read_excel(file_path):
    wb = openpyxl.load_workbook(file_path)
    sheet = wb.active
    data = []
    for row in sheet.iter_rows(values_only=True):
        data.append(row)
    return data


# 将 Excel 数据转换为 Word 表格
def write_to_word_with_table(data, output_path):
    doc = Document()

    # 创建一个表格，行数和列数与 Excel 文件相同
    table = doc.add_table(rows=len(data), cols=len(data[0]))

    # 填充表格
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            cell_text = str(cell) if cell is not None else ''
            table.cell(i, j).text = cell_text

            # 设置表格样式
            table.cell(i, j).paragraphs[0].runs[0].font.size = Pt(12)  # 字体大小
            table.cell(i, j).paragraphs[0].alignment = 1  # 居中对齐

            # 设置表格单元格的背景颜色（例如：灰色背景）
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            table.cell(i, j)._element.get_or_add_tcPr().append(shading_elm)

    doc.save(output_path)


# 主函数
def excel_to_word_with_table_format(excel_path, word_path):
    data = read_excel(excel_path)
    write_to_word_with_table(data, word_path)
    print(f"Excel 数据已成功转换为 Word 文档，并保留了表格格式：{word_path}")


# 示例：将 Excel 转换为带表格的 Word 文件
excel_to_word_with_table_format('C:/Users/win11/Desktop/123.xlsx', 'C:/Users/win11/Desktop/new.docx')
